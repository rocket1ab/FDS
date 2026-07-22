#!/usr/bin/env python3
"""Update section 6 DDA tables in the retained aircraft assessment DOCX."""

from __future__ import annotations

import argparse
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ANGLE_ROWS = [
    ["el=15°", "1,912", "757.1", "735", "0.971", "48.5"],
    ["el=30°", "1,434", "757.1", "647", "0.855", "42.7"],
    ["el=45°", "1,344", "757.1", "529", "0.699", "34.9"],
]

MATERIAL_ROWS = [
    ["RADM 雷达罩\n玻璃纤维复合材料", "386", "735", "48.5", "97.1%", "直接受照"],
    ["WINS 舷窗\nPMMA有机玻璃", "213", "677", "44.7", "89.4%", "直接受照"],
    ["BED 床垫\n尼龙织物表层", "99", "441", "29.1", "58.2%", "直接受照"],
    ["CURT 窗帘\n尼龙织物", "267", "618", "40.8", "81.6%", "直接受照"],
    ["U4 仪表设备\n原U04材料参数", "0", "0", "0", "0%", "遮挡/二次受热"],
    ["SEAT 座椅\n聚氨酯泡沫", "92", "618", "40.8", "81.6%", "直接受照"],
    ["AL2024 蒙皮\n2024铝合金", "674", "735", "48.5", "97.1%", "直接受照"],
    ["AL5052 风管\n5052铝合金", "0", "0", "0", "0%", "遮挡/二次受热"],
    ["AL7075 隔框\n7075铝合金", "121", "735", "48.5", "97.1%", "直接受照"],
    ["O2TANK 氧气瓶\n7075铝合金", "10", "529", "34.9", "69.9%", "局部受照"],
    ["H1 导航子系统\n6061铝合金3 mm", "5", "735", "48.5", "97.1%", "局部受照"],
    ["H2 任务子系统\n6061铝合金3 mm", "37", "735", "48.5", "97.1%", "局部受照"],
    ["H3 显示子系统\n6061铝合金3 mm", "0", "0", "0", "0%", "遮挡/二次受热"],
    ["H4 通信子系统\n6061铝合金3 mm", "4", "118", "7.8", "15.6%", "弱局部受照"],
    ["H5 电池\n6061铝外壳3 mm", "4", "618", "40.8", "81.6%", "局部受照"],
    ["H6 电力传输\nPVC 1 mm", "0", "0", "0", "0%", "遮挡/二次受热"],
    ["H7 操纵子系统\nCR橡胶2 mm", "0", "0", "0", "0%", "遮挡/二次受热"],
]


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def fill_table(table, headers: list[str], rows: list[list[str]], font_size: float) -> None:
    while len(table.columns) < len(headers):
        table.add_column(Inches(0.85))
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    while len(table.rows) > len(rows) + 1:
        table._tbl.remove(table.rows[-1]._tr)
    for col, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[col], text, bold=True, size=font_size)
    repeat_header(table.rows[0])
    for row_index, values in enumerate(rows, start=1):
        prevent_row_split(table.rows[row_index])
        for col, text in enumerate(values):
            set_cell_text(table.rows[row_index].cells[col], text, size=font_size)


def insert_paragraph_before_table(document: Document, table, text: str, style: str = "Body Text"):
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)
    table._tbl.addprevious(paragraph._p)
    return paragraph


def style_added_paragraph(paragraph, size: float = 9.5) -> None:
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()
    source = args.docx.resolve()
    document = Document(source)

    angle_table = next(
        table for table in document.tables
        if table.rows and "角度" in table.rows[0].cells[0].text
        and any("DDA受照体素面数" in cell.text for cell in table.rows[0].cells)
    )
    material_table = next(
        table for table in document.tables
        if table.rows and "材料/设备组" in table.rows[0].cells[0].text
        and any("EXTERNAL_FLUX" in cell.text for cell in table.rows[0].cells)
    )
    fill_table(
        angle_table,
        ["俯仰角", "DDA受照\n体素面数", "入射平面峰值\n(kW/m²)", "最大局部峰值\n(kW/m²)",
         "最大几何\n传递系数", "最大局部光冲量\n(J/cm²)"],
        ANGLE_ROWS,
        7.2,
    )
    fill_table(
        material_table,
        ["材料/设备组", "DDA受照\n面数", "最大局部峰值\n(kW/m²)", "最大局部光冲量\n(J/cm²)",
         "标称Q\n传递比例", "当前角度\n受照状态"],
        MATERIAL_ROWS,
        7.0,
    )

    heading_62 = next(p for p in document.paragraphs if p.text.strip() == "6.2 各材料受照分析")
    normalization_text = (
        "Q=50 J/cm²、W=100 kt时，NUCLEAR_RAMP积分为0.660398 s，按E₀=10Q/∫F(t)dt得到入射平面峰值"
        "757.1 kW/m²。最大几何传递系数为最大局部峰值与入射平面峰值之比；对应el=15°、30°和45°的"
        "最大局部积分光冲量分别为48.5、42.7和34.9 J/cm²。新版结果不再与旧E₀=1287 kW/m²基准缩放混用。"
    )
    normalization = next((p for p in document.paragraphs if "NUCLEAR_RAMP积分" in p.text), None)
    if normalization is None:
        normalization = heading_62.insert_paragraph_before(normalization_text, style="Body Text")
    else:
        normalization.text = normalization_text
    style_added_paragraph(normalization)

    intro_text = (
        "以下结果对应az=270°、el=15°、Q=50 J/cm²。局部积分光冲量由Q归一化局部峰值与同一"
        "NUCLEAR_RAMP积分计算；受照面数表示DDA判定具有直达视线且写入EXTERNAL_FLUX的0.1 m体素面数量。"
    )
    intro = next((p for p in document.paragraphs if p.text.startswith("以下结果对应az=270°")), None)
    if intro is None:
        intro = insert_paragraph_before_table(document, material_table, intro_text)
    else:
        intro.text = intro_text
    style_added_paragraph(intro)

    old_note = next((p for p in document.paragraphs if
                     ("AL_FLUX_MAP" in p.text and "vendor" in p.text) or
                     "旧版“铝表面未赋热流”问题已修复" in p.text), None)
    if old_note is not None:
        old_note.text = (
            "说明：新版体素流程已为AL2024、AL7075、H1-H5等铝合金表面生成EXTERNAL_FLUX变体，"
            "旧版“铝表面未赋热流”问题已修复。表中0受照面表示设备在az=270°、el=15°下被几何遮挡、"
            "没有直达外部辐射；其材料属性仍保留，并可通过舱内火焰、热烟气和二次辐射升温。"
        )
        style_added_paragraph(old_note, 9.0)

    figure_dir = source.parent
    figures = [
        (angle_table, figure_dir / "section6_angle_exposure_fluence.png",
         "图6.3 不同俯仰角的DDA受照规模与最大局部积分光冲量。"),
        (material_table, figure_dir / "section6_material_faces_peakflux.png",
         "图6.4 17组材料/设备的受照面数与最大局部峰值热流。"),
    ]
    for table, image_path, caption in figures:
        if not image_path.exists():
            continue
        picture_paragraph = document.add_paragraph()
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.add_run().add_picture(str(image_path), width=Inches(6.25))
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.add_run(caption)
        caption_run.italic = True
        caption_run.font.name = "宋体"
        caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        caption_run.font.size = Pt(8.5)
        table._tbl.addnext(caption_paragraph._p)
        table._tbl.addnext(picture_paragraph._p)

    for paragraph in document.paragraphs:
        if old_note is not None and paragraph is old_note:
            continue
        if "铝 EXTERNAL_FLUX 基座未生成" in paragraph.text:
            paragraph.text = (
                "铝合金外部热流赋值已修复：新版DDA体素流程已为AL2024、AL7075及H1-H5生成"
                "EXTERNAL_FLUX变体；后续角度比较均应采用Q归一化后的新版结果。"
            )
            style_added_paragraph(paragraph)

    temp = source.with_name(source.stem + ".section6.tmp.docx")
    document.save(temp)
    shutil.copy2(temp, source)
    temp.unlink()
    print(source)


if __name__ == "__main__":
    main()
